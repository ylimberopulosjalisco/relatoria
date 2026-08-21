
import io
import re
import hashlib
import textwrap
from collections import Counter
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st

try:
    from supabase import create_client
except Exception:
    create_client = None

try:
    from PIL import Image as PILImage, ImageDraw, ImageFont
except Exception:
    PILImage = ImageDraw = ImageFont = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, PageBreak, KeepTogether
    )
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except Exception:
    colors = None


# ==========================================================
# RELATORÍA COINVIERTE — OTRA BARRERA Y OTRO ACTOR SIEMPRE HABILITADOS
# ==========================================================

st.set_page_config(
    page_title="Relatoría | Economía Circular Jalisco",
    page_icon="🌿",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------- Paleta ----------
INK = "#173B2E"
INK_2 = "#285843"
SAGE = "#DCEBCB"
SAGE_2 = "#EEF5E8"
MINT = "#EAF4F1"
AQUA = "#D9EFF1"
AQUA_STRONG = "#3C8791"
SAND = "#F4E6C8"
CREAM = "#FBFAF5"
WHITE = "#FFFFFF"
BORDER = "#DDE5DA"
TEXT = "#2D3C33"
MUTED = "#6E7B73"
GREEN = "#2F6E4A"
GREEN_DARK = "#24583B"
SHADOW = "rgba(38, 72, 51, 0.08)"

st.markdown(
    f"""
    <style>
    /* ---------- Ocultar chrome superior de Streamlit ---------- */
    header[data-testid="stHeader"] {{
        display: none !important;
        height: 0 !important;
        min-height: 0 !important;
    }}
    [data-testid="stToolbar"],
    [data-testid="stDecoration"],
    [data-testid="stStatusWidget"] {{
        display: none !important;
    }}

    /* ---------- Página ---------- */
    .stApp {{
        background:
            radial-gradient(circle at 75% 0%, rgba(217,239,241,.55), transparent 28%),
            radial-gradient(circle at 8% 100%, rgba(220,235,203,.48), transparent 24%),
            linear-gradient(180deg, #FFFDF9 0%, #FBFCF8 100%);
        color: {TEXT};
    }}
    .block-container {{
        max-width: 1460px;
        padding-top: 0.75rem !important;
        padding-bottom: 2.5rem;
    }}

    h1, h2, h3 {{
        color: {INK};
        letter-spacing: -0.02em;
    }}

    /* ---------- Sidebar ---------- */
    [data-testid="stSidebar"] {{
        background:
            linear-gradient(180deg, rgba(255,253,248,.98), rgba(249,250,244,.98));
        border-right: 1px solid {BORDER};
        position: relative;
        overflow: hidden;
        min-width: 340px !important;
        width: 340px !important;
    }}
    [data-testid="stSidebar"] > div:first-child {{
        width: 340px !important;
    }}
    [data-testid="stSidebar"]::after {{
        content: "";
        position: absolute;
        left: -70px;
        bottom: -90px;
        width: 330px;
        height: 210px;
        background:
            radial-gradient(ellipse at 60% 80%, rgba(60,135,145,.18) 0 32%, transparent 33%),
            radial-gradient(ellipse at 28% 78%, rgba(220,235,203,.95) 0 36%, transparent 37%),
            radial-gradient(ellipse at 5% 80%, rgba(244,230,200,.95) 0 40%, transparent 41%);
        border-radius: 45% 55% 0 0 / 80% 80% 0 0;
        pointer-events: none;
        z-index: 0;
    }}
    [data-testid="stSidebar"] {{
        min-width: 410px !important;
        width: 410px !important;
    }}
    [data-testid="stSidebar"] > div {{
        position: relative;
        z-index: 1;
        width: 410px !important;
    }}

    /* Selectores largos: mostrar el texto completo, sin cortarlo con puntos suspensivos */
    [data-baseweb="select"] > div {{
        min-height: 42px !important;
        height: auto !important;
    }}
    [data-baseweb="select"] [data-testid="stMarkdownContainer"],
    [data-baseweb="select"] span,
    div[role="option"] {{
        white-space: normal !important;
        overflow: visible !important;
        text-overflow: clip !important;
        line-height: 1.25 !important;
    }}
    div[role="option"] {{
        min-height: 42px !important;
        height: auto !important;
        align-items: flex-start !important;
        padding-top: 10px !important;
        padding-bottom: 10px !important;
    }}
    [data-testid="stSidebar"] h2 {{
        color: {INK};
        font-weight: 800;
        margin-bottom: .9rem;
    }}
    [data-testid="stSidebar"] label {{
        color: {TEXT} !important;
        font-weight: 650 !important;
    }}
    [data-testid="stSidebar"] .stSelectbox > div > div,
    [data-testid="stSidebar"] .stTextInput input {{
        background: rgba(255,255,255,.95) !important;
        border-color: {BORDER} !important;
        border-radius: 12px !important;
    }}

    /* ---------- Hero ---------- */
    .hero-shell {{
        position: relative;
        background:
            radial-gradient(circle at 88% 5%, rgba(255,255,255,.55), transparent 18%),
            linear-gradient(120deg, #FFFDF7 0%, #FFF9ED 48%, #EFF7F5 100%);
        border: 1px solid rgba(221,229,218,.65);
        border-radius: 28px;
        min-height: 205px;
        padding: 34px 42px 30px 42px;
        overflow: hidden;
        box-shadow: 0 10px 28px {SHADOW};
    }}
    .hero-shell::before {{
        content: "";
        position: absolute;
        right: -120px;
        bottom: -95px;
        width: 520px;
        height: 255px;
        border-radius: 55% 45% 0 0 / 70% 70% 0 0;
        background:
            radial-gradient(ellipse at 65% 80%, rgba(60,135,145,.14) 0 42%, transparent 43%),
            radial-gradient(ellipse at 38% 80%, rgba(220,235,203,.45) 0 49%, transparent 50%);
    }}
    .hero-shell::after {{
        content: "";
        position: absolute;
        right: 76px;
        top: 35px;
        width: 105px;
        height: 105px;
        border: 13px solid rgba(47,110,74,.13);
        border-radius: 50%;
        clip-path: polygon(0 0,100% 0,100% 38%,60% 38%,60% 100%,0 100%);
        transform: rotate(15deg);
        opacity: .85;
    }}
    .hero-title {{
        position: relative;
        z-index: 2;
        color: {INK};
        font-size: 2.45rem;
        font-weight: 850;
        line-height: 1.05;
        margin-bottom: .55rem;
        max-width: 850px;
    }}
    .hero-sub {{
        position: relative;
        z-index: 2;
        color: {TEXT};
        font-size: 1rem;
        max-width: 820px;
        line-height: 1.55;
    }}
    .pill-row {{
        position: relative;
        z-index: 2;
        margin-top: 1.05rem;
    }}
    .soft-pill {{
        display: inline-block;
        margin: 0 .42rem .35rem 0;
        padding: .45rem .8rem;
        border-radius: 999px;
        font-size: .83rem;
        font-weight: 650;
        color: {INK};
    }}
    .pill-sage {{ background: {SAGE}; }}
    .pill-sand {{ background: {SAND}; }}
    .pill-aqua {{ background: {AQUA}; }}

    /* ---------- Franja institucional de logos ---------- */
    .brand-separator {{
        width: 1px;
        height: 48px;
        background: {BORDER};
        margin: 0 auto;
    }}

    /* ---------- Tarjetas ---------- */
    .info-card {{
        background: rgba(255,255,255,.96);
        border: 1px solid {BORDER};
        border-radius: 18px;
        padding: 16px 18px;
        min-height: 98px;
        box-shadow: 0 6px 18px {SHADOW};
        display: flex;
        align-items: center;
        gap: 13px;
    }}
    .info-icon {{
        width: 48px;
        height: 48px;
        border-radius: 50%;
        display:flex;
        align-items:center;
        justify-content:center;
        font-size: 1.35rem;
        flex: 0 0 auto;
    }}
    .icon-sage {{ background: {SAGE}; }}
    .icon-sand {{ background: {SAND}; }}
    .icon-aqua {{ background: {AQUA}; }}
    .info-label {{
        font-size: .72rem;
        letter-spacing: .05em;
        color: {MUTED};
        font-weight: 800;
        text-transform: uppercase;
        margin-bottom: .2rem;
    }}
    .info-value {{
        color: {INK};
        font-size: 1.03rem;
        font-weight: 800;
        line-height: 1.25;
    }}

    /* ---------- Tabs ---------- */
    .stTabs [data-baseweb="tab-list"] {{
        gap: 0.1rem;
        border-bottom: 1px solid {BORDER};
    }}
    .stTabs [data-baseweb="tab"] {{
        height: auto;
        padding: .55rem .95rem;
        background: transparent;
        border-radius: 10px 10px 0 0;
        color: {TEXT};
        font-weight: 650;
    }}
    .stTabs [aria-selected="true"] {{
        color: {INK} !important;
        background: rgba(238,245,232,.7) !important;
        border-bottom: 3px solid {GREEN} !important;
    }}

    /* ---------- Formulario ---------- */
    .rule-note {{
        background: linear-gradient(90deg, rgba(220,235,203,.75), rgba(238,245,232,.72));
        border-radius: 14px;
        padding: 12px 16px;
        margin-bottom: 14px;
        border: 1px solid rgba(221,229,218,.85);
        color: {TEXT};
    }}

    [data-testid="stForm"] {{
        background: rgba(255,255,255,.90);
        border: 1px solid {BORDER};
        border-radius: 20px;
        padding: 16px 18px 18px 18px;
        box-shadow: 0 7px 20px {SHADOW};
    }}
    .stTextArea textarea,
    .stTextInput input,
    .stSelectbox > div > div {{
        background: #FFFDF9 !important;
        border-color: {BORDER} !important;
        border-radius: 12px !important;
    }}

    /* ---------- Botones ---------- */
    div.stButton > button,
    div[data-testid="stFormSubmitButton"] > button {{
        border-radius: 13px !important;
        min-height: 46px;
        font-weight: 750 !important;
    }}
    div[data-testid="stFormSubmitButton"] > button[kind="primary"] {{
        background: linear-gradient(120deg, {GREEN}, {INK_2}) !important;
        color: white !important;
        border: none !important;
    }}
    div[data-testid="stFormSubmitButton"] > button[kind="primary"]:hover {{
        background: linear-gradient(120deg, {GREEN_DARK}, {INK}) !important;
    }}

    /* ---------- Sidebar / selectores largos ---------- */
    section[data-testid="stSidebar"] {{
        width: 390px !important;
        min-width: 390px !important;
    }}
    section[data-testid="stSidebar"] > div {{
        width: 390px !important;
    }}

    /* El valor seleccionado debe poder leerse completo */
    section[data-testid="stSidebar"] div[data-baseweb="select"] > div {{
        min-height: 44px !important;
        height: auto !important;
    }}
    section[data-testid="stSidebar"] div[data-baseweb="select"] span {{
        white-space: normal !important;
        overflow: visible !important;
        text-overflow: clip !important;
        line-height: 1.25 !important;
    }}

    /* El menú desplegable de BaseWeb vive fuera del sidebar.
       Ensancharlo y permitir salto de línea evita que se corte. */
    div[data-baseweb="popover"] {{
        min-width: 460px !important;
        width: max-content !important;
        max-width: min(620px, 80vw) !important;
    }}
    div[data-baseweb="popover"] [role="listbox"] {{
        min-width: 460px !important;
        width: max-content !important;
        max-width: min(620px, 80vw) !important;
    }}
    div[data-baseweb="popover"] [role="option"] {{
        white-space: normal !important;
        overflow: visible !important;
        text-overflow: clip !important;
        line-height: 1.3 !important;
        height: auto !important;
        min-height: 42px !important;
        padding-top: 10px !important;
        padding-bottom: 10px !important;
    }}
    div[data-baseweb="popover"] [role="option"] > div,
    div[data-baseweb="popover"] [role="option"] span {{
        white-space: normal !important;
        overflow: visible !important;
        text-overflow: clip !important;
        max-width: none !important;
    }}

    /* ---------- Sidebar notes ---------- */
    .sidebar-note {{
        background: linear-gradient(120deg, rgba(234,244,241,.96), rgba(255,255,255,.92));
        border: 1px solid {BORDER};
        border-radius: 14px;
        padding: 13px 14px;
        color: {MUTED};
        line-height: 1.45;
        margin-bottom: .75rem;
    }}
    .sidebar-db {{
        background: linear-gradient(120deg, rgba(220,235,203,.92), rgba(238,245,232,.95));
        border: 1px solid {BORDER};
        border-radius: 14px;
        padding: 13px 14px;
        color: {INK};
        font-weight: 750;
    }}

    /* ---------- Métricas / tablas ---------- */
    div[data-testid="stMetric"] {{
        background: rgba(255,255,255,.96);
        border: 1px solid {BORDER};
        border-radius: 17px;
        padding: 10px 14px;
        box-shadow: 0 6px 18px {SHADOW};
    }}

    /* ---------- Footer ---------- */
    .app-footer {{
        text-align:center;
        color:{MUTED};
        font-size:.82rem;
        margin-top:.8rem;
    }}

    /* ---------- Responsive ---------- */
    @media (max-width: 900px) {{
        .hero-title {{ font-size: 1.75rem; }}
        .hero-shell {{ border-radius: 24px; padding: 22px; }}
        .logo-panel {{ min-height: auto; }}
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

BASE_DIR = Path(__file__).resolve().parent
ASSETS_DIR = BASE_DIR / "assets"

def buscar_logo(preferidos=None, terminos=None):
    preferidos = preferidos or []
    terminos = terminos or []
    if not ASSETS_DIR.exists():
        return None

    # 1) Buscar por nombres preferidos exactos
    for nombre in preferidos:
        p = ASSETS_DIR / nombre
        if p.exists():
            return p

    # 2) Buscar por términos en nombre de archivo
    for candidate in sorted(ASSETS_DIR.iterdir()):
        if candidate.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
            nombre = candidate.name.lower()
            if all(t in nombre for t in terminos):
                return candidate

    # 3) Buscar por cualquier término
    for candidate in sorted(ASSETS_DIR.iterdir()):
        if candidate.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
            nombre = candidate.name.lower()
            if any(t in nombre for t in terminos):
                return candidate

    return None

LOGO_COINVIERTE = buscar_logo(
    preferidos=[
        "logo_coinvierte.png",
        "logo_coinvierte.jpg",
        "logo_coinvierte.jpeg",
        "logo_coinvierte.webp",
    ],
    terminos=["coinvierte"]
)

LOGO_TEC = buscar_logo(
    preferidos=[
        "logo_tec_monterrey.png",
        "logo_tec_monterrey.jpg",
        "logo_tec_monterrey.jpeg",
        "logo_tec.png",
        "tec_monterrey.png",
        "tecnologico_de_monterrey.png",
        "tecnologico_de_monterrey.jpg",
    ],
    terminos=["tec", "monterrey"]
)

# Fallback por si el logo del Tec sólo contiene "tec" o "tecnologico"
if LOGO_TEC is None:
    LOGO_TEC = buscar_logo(
        terminos=["tecnologico"]
    ) or buscar_logo(terminos=["tec"])

DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(exist_ok=True)
LOCAL_CSV = DATA_DIR / "relatoria_hallazgos.csv"

MESAS = [
    "Estrategia, gobernanza y capacidades",
    "Operación circular, recursos y cadenas de valor",
    "Innovación, tecnología, inversión y financiamiento",
    "Entorno habilitante, regulación y colaboración",
]

PREGUNTAS_POR_MESA = {
    "Estrategia, gobernanza y capacidades": [
        "Detonadora. Pensando en su organización, ¿en qué medida la economía circular ya forma parte de la manera en que toman decisiones y operan, y en qué medida sigue siendo algo incipiente?",
        "P1. ¿Su organización tiene hoy alguna práctica o iniciativa relacionada con economía circular, aunque sea incipiente o no la llamen así?",
        "P2. Cuando aparece una oportunidad de circularidad, ¿quién puede realmente impulsarla o detenerla dentro de la organización?",
        "P3. ¿Qué les hace falta dentro de la organización para poder avanzar más rápido?",
        "P4. ¿Qué información tendría que llegar a su escritorio para que ustedes dijeran: “sí, vale la pena hacer esta inversión”?",
        "Cierre. Si pudieran fortalecer una sola capacidad interna durante los próximos dos años, ¿cuál tendría mayor impacto?",
    ],
    "Operación circular, recursos y cadenas de valor": [
        "Detonadora. Si recorriéramos hoy su operación completa, ¿en qué parte encontraríamos la mayor pérdida de valor: materiales, residuos, agua, energía, empaques, inventarios o logística?",
        "P1. ¿Qué flujo de recursos consideran que tendría mayor potencial de mejora, aun si todavía no han iniciado acciones para atenderlo?",
        "P2. ¿Qué les impide aprovechar ese material, reducir ese consumo o cerrar ese ciclo actualmente?",
        "P3. ¿Identifican algún residuo o subproducto que potencialmente pudiera aprovechar otra empresa, o algún material externo que ustedes pudieran aprovechar, aunque hoy ese intercambio no exista?",
        "P4. ¿Qué tendría que cambiar en su cadena de proveedores o clientes para que ustedes pudieran ser más circulares?",
        "Cierre. ¿Cuál sería la oportunidad operativa más tangible que podríamos comenzar a resolver en Jalisco?",
    ],
    "Innovación, tecnología, inversión y financiamiento": [
        "Detonadora. ¿Qué problema de su empresa saben que podría resolverse con tecnología o innovación, pero todavía no han podido resolver?",
        "P1. ¿Qué tecnologías o soluciones conocen, han explorado o creen que podrían ser relevantes para avanzar hacia una economía más circular? Si ninguna está hoy en el radar, ¿por qué?",
        "P2. ¿Qué hace que una solución técnicamente viable no termine implementándose?",
        "P3. ¿Qué tendría que demostrar o resolver un proyecto para que su empresa considerara invertir recursos propios, incluso si hoy no existe disposición de inversión?",
        "P4. ¿Qué parte del riesgo tendría que compartir alguien más para que el proyecto ocurriera?",
        "P5. Si el gobierno quisiera ayudar a destrabar estas inversiones, ¿qué sería más útil y por qué?",
        "Cierre. ¿Qué inversión circular creen que hoy no está ocurriendo en Jalisco pero podría ocurrir en los próximos dos o tres años si se elimina la barrera correcta?",
    ],
    "Entorno habilitante, regulación y colaboración": [
        "Detonadora. Pensando fuera de su empresa, ¿cuál es hoy el principal obstáculo del entorno para avanzar hacia una economía circular?",
        "P1. ¿Hay alguna regulación, trámite o falta de claridad regulatoria que esté dificultando una solución circular?",
        "P2. ¿Qué tendría que cambiar en el mercado para que las soluciones circulares fueran más competitivas?",
        "P3. ¿Qué problema no puede resolver una empresa sola y requeriría colaboración entre varias organizaciones?",
        "P4. ¿Qué debería hacer el Gobierno de Jalisco para acelerar la economía circular que hoy no está haciendo, o debería hacer de manera diferente?",
        "Cierre. Si pudiéramos modificar una sola condición del ecosistema de Jalisco durante los próximos tres años, ¿cuál tendría mayor efecto?",
    ],
}

PREGUNTAS_LIDERES_GREMIALES = {
    "Estrategia, gobernanza y capacidades": [
        "Detonadora. Desde su posición gremial, ¿qué tan incorporada está hoy la economía circular en la agenda de las empresas de su sector?",
        "P1. ¿Qué tipo de empresas están avanzando más y cuáles se están quedando atrás? ¿Por qué?",
        "P2. ¿Cuáles son hoy las principales brechas de capacidades del sector: conocimiento técnico, información, talento, liderazgo, indicadores u otras?",
        "P3. ¿Qué tendría que cambiar para que más empresas incorporen la circularidad como una decisión estratégica y no sólo como cumplimiento o una iniciativa aislada?",
        "P4. ¿Qué papel podrían jugar las cámaras y asociaciones empresariales para acelerar esa transición?",
        "Cierre. Si pudiera fortalecerse una sola capacidad sectorial en los próximos dos años, ¿cuál tendría mayor efecto?",
    ],
    "Operación circular, recursos y cadenas de valor": [
        "Detonadora. Desde una perspectiva sectorial, ¿dónde están hoy las mayores pérdidas de valor en materiales, agua, energía, residuos o logística?",
        "P1. ¿Qué flujos de residuos o subproductos tienen mayor potencial de aprovechamiento entre empresas del sector o entre sectores?",
        "P2. ¿Qué está impidiendo que esos intercambios ocurran hoy de manera sistemática?",
        "P3. ¿Qué problemas de escala, logística, calidad, trazabilidad o regulación aparecen de manera recurrente entre sus afiliados?",
        "P4. ¿Qué tipo de infraestructura o mecanismo compartido podría generar mayor impacto: centros de acopio, plataformas de intercambio, logística conjunta, servicios especializados u otro?",
        "Cierre. ¿Cuál sería la oportunidad de simbiosis o articulación empresarial más viable para comenzar a trabajar en Jalisco?",
    ],
    "Innovación, tecnología, inversión y financiamiento": [
        "Detonadora. ¿Qué tecnologías o soluciones circulares considera que tienen mayor potencial de adopción en las empresas de su sector durante los próximos años?",
        "P1. ¿Qué tecnologías están maduras, pero siguen sin escalar entre las empresas? ¿Qué las frena?",
        "P2. ¿Qué tipo de riesgo perciben las empresas: tecnológico, financiero, comercial, regulatorio o de implementación?",
        "P3. ¿Qué características tendría que tener un instrumento de apoyo para que realmente movilizara inversión privada?",
        "P4. ¿Qué hace más sentido para su sector: coinversión, crédito, garantías, asistencia técnica, pilotos, compras agregadas, vinculación con proveedores u otro mecanismo?",
        "P5. ¿Qué tipo de proyectos podrían estructurarse de manera colectiva o sectorial, en lugar de empresa por empresa?",
        "Cierre. Si Jalisco pudiera habilitar una sola línea de inversión circular para su sector, ¿cuál tendría mayor capacidad de movilizar capital privado?",
    ],
    "Entorno habilitante, regulación y colaboración": [
        "Detonadora. ¿Cuáles son hoy las principales condiciones del entorno que frenan la economía circular en su sector?",
        "P1. ¿Qué regulación, trámite, falta de claridad o ausencia de estándares está generando mayores obstáculos?",
        "P2. ¿Qué incentivos o señales de mercado podrían acelerar la adopción de prácticas circulares?",
        "P3. ¿Qué problemas requieren coordinación entre gobierno, empresas, academia, banca y organizaciones empresariales?",
        "P4. ¿Qué debería hacer el Gobierno de Jalisco que hoy no está haciendo para acelerar la transición circular?",
        "P5. ¿Qué podría hacer mejor el propio sector organizado —cámaras, asociaciones y organismos empresariales— para facilitar esa transición?",
        "Cierre. Si pudiera cambiar una sola condición de política pública o del ecosistema empresarial en Jalisco, ¿cuál tendría mayor impacto?",
    ],
}


GRUPOS = [
    "Sector primario",
    "Sector secundario",
    "Sector terciario",
    "Líderes gremiales",
]

BARRERAS = [
    "Sin clasificar",
    "Técnica",
    "Financiera",
    "Regulatoria",
    "Mercado",
    "Información / datos",
    "Capacidades",
    "Proveedores / infraestructura",
    "Coordinación",
    "Prioridad / gobernanza",
    "Otra",
]

ACTORES = [
    "Sin definir",
    "Empresa",
    "Gobierno estatal",
    "Gobierno municipal",
    "Cámara / organismo empresarial",
    "Academia / centro tecnológico",
    "Banca / financiador",
    "Proveedor",
    "Varias empresas",
    "Otro",
]


TIPOS_HALLAZGO_GREMIAL = [
    "Barrera sistémica",
    "Brecha de capacidades",
    "Regulación / trámite",
    "Mercado / demanda",
    "Infraestructura",
    "Financiamiento",
    "Tecnología / innovación",
    "Coordinación entre actores",
    "Oportunidad sectorial",
    "Otro",
]

AFECTACION_GREMIAL = [
    "PyMEs",
    "Grandes empresas",
    "Todo el sector",
    "Ciertos subsectores",
    "No se sabe",
    "Otro",
]

ACTORES_GREMIALES = [
    "Gobierno estatal",
    "Gobierno municipal",
    "Gobierno federal",
    "Cámaras / asociaciones",
    "Empresas",
    "Academia",
    "Banca / financiadores",
    "Proveedores",
    "Otro",
]

INSTRUMENTOS_GREMIALES = [
    "Regulación / simplificación",
    "Incentivo",
    "Coinversión",
    "Crédito / garantía",
    "Asistencia técnica",
    "Capacitación",
    "Infraestructura compartida",
    "Plataforma / coordinación",
    "Piloto",
    "Otro",
]

SECTORIALIDAD = [
    "No se sabe",
    "Sí, parece sectorial",
    "No, parece particular de la empresa",
]

PRIORIDADES = ["Sin clasificar", "Alta", "Media", "Baja"]


# ==========================================================
# Persistencia
# ==========================================================
@st.cache_resource
def get_supabase():
    if create_client is None:
        return None
    try:
        url = st.secrets.get("SUPABASE_URL", "")
        key = st.secrets.get(
            "SUPABASE_PUBLISHABLE_KEY",
            st.secrets.get("SUPABASE_KEY", "")
        )
        if url and key:
            return create_client(url, key)
    except Exception:
        pass
    return None


supabase = get_supabase()


def load_data():
    cols = [
        "id", "fecha_hora", "mesa", "grupo", "ronda", "relator", "moderador",
        "tipo_pregunta", "pregunta_referencia",
        "hallazgo", "barrera", "barrera_otra", "ejemplo", "actor", "actor_otro",
        "apoyo_solucion", "sectorialidad", "prioridad", "frase_clave", "notas",
        "tipo_hallazgo_gremial", "afectacion_gremial", "instrumento_gremial"
    ]

    if supabase:
        try:
            rows = (
                supabase.table("relatoria_hallazgos")
                .select("*")
                .order("created_at", desc=False)
                .execute()
                .data or []
            )
            df = pd.DataFrame(rows)
            if not df.empty:
                if "created_at" in df.columns and "fecha_hora" not in df.columns:
                    df["fecha_hora"] = df["created_at"]
                return df
        except Exception as e:
            st.warning(
                f"No se pudo leer Supabase. Se usará almacenamiento local. Detalle: {e}"
            )

    if LOCAL_CSV.exists():
        try:
            return pd.read_csv(LOCAL_CSV)
        except Exception:
            pass

    return pd.DataFrame(columns=cols)


def save_record(record):
    if supabase:
        payload = {k: v for k, v in record.items() if k != "id"}
        payload["created_at"] = datetime.now().isoformat(timespec="seconds")
        try:
            supabase.table("relatoria_hallazgos").insert(payload).execute()
            return True, "Guardado en Supabase."
        except Exception as e:
            return False, f"No se pudo guardar en Supabase: {e}"

    df = load_data()

    if df.empty:
        new_id = 1
    else:
        ids = pd.to_numeric(df.get("id", pd.Series(dtype=float)), errors="coerce")
        new_id = int(ids.max()) + 1 if ids.notna().any() else 1

    record = dict(record)
    record["id"] = new_id
    df = pd.concat([df, pd.DataFrame([record])], ignore_index=True)
    df.to_csv(LOCAL_CSV, index=False)
    return True, "Guardado localmente."


def delete_record(record_id):
    if supabase:
        try:
            supabase.table("relatoria_hallazgos").delete().eq(
                "id", int(record_id)
            ).execute()
            return True
        except Exception as e:
            st.error(f"No se pudo borrar: {e}")
            return False

    df = load_data()
    if "id" in df.columns:
        df = df[df["id"].astype(str) != str(record_id)]
        df.to_csv(LOCAL_CSV, index=False)
        return True
    return False


def to_excel_bytes(df):
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Hallazgos")
        if not df.empty:
            resumen = (
                df.groupby(["mesa", "grupo", "barrera"], dropna=False)
                .size()
                .reset_index(name="n")
                .sort_values(
                    ["mesa", "grupo", "n"],
                    ascending=[True, True, False]
                )
            )
            resumen.to_excel(writer, index=False, sheet_name="Resumen")
    return out.getvalue()



# ==========================================================
# Síntesis cualitativa y exportación
# ==========================================================
THEME_KEYWORDS = {
    "Financiamiento e inversión": [
        "financ", "invers", "crédito", "credito", "garant", "costo", "presupuesto", "capital"
    ],
    "Capacidades y conocimiento": [
        "capacit", "conocimiento", "talento", "personal", "habilidad", "formación", "formacion", "asistencia técnica", "asistencia tecnica"
    ],
    "Regulación y trámites": [
        "regul", "norma", "trámite", "tramite", "permiso", "legal", "claridad regulatoria"
    ],
    "Coordinación y colaboración": [
        "coordin", "colabor", "alianza", "vincul", "articul", "grem", "cámara", "camara", "ecosistema"
    ],
    "Tecnología e innovación": [
        "tecnolog", "innov", "digital", "automat", "plataforma", "trazab", "equipo", "maquinaria"
    ],
    "Mercado y demanda": [
        "mercado", "cliente", "demanda", "precio", "compet", "comercial", "venta", "comprador"
    ],
    "Proveedores y cadena de valor": [
        "proveedor", "cadena", "suministro", "logíst", "logist", "insumo", "cliente"
    ],
    "Información y medición": [
        "información", "informacion", "dato", "medición", "medicion", "indicador", "evidencia", "monitoreo"
    ],
    "Infraestructura": [
        "infraestructura", "instalación", "instalacion", "centro", "acopio", "tratamiento", "planta"
    ],
    "Materiales y residuos": [
        "residuo", "material", "subproducto", "recicl", "reuso", "reutil", "merma", "desecho", "empaque"
    ],
    "Agua y energía": [
        "agua", "energ", "electric", "combustible", "renovable", "consumo"
    ],
}

ACTION_BY_THEME = {
    "Financiamiento e inversión": "estructurar mecanismos de coinversión, crédito o garantías que reduzcan el riesgo de implementación",
    "Capacidades y conocimiento": "fortalecer asistencia técnica, capacitación y capacidades internas para convertir oportunidades en proyectos ejecutables",
    "Regulación y trámites": "identificar cuellos de botella regulatorios y generar mayor claridad sobre permisos, criterios y rutas de cumplimiento",
    "Coordinación y colaboración": "crear mecanismos de articulación entre empresas, gobierno, academia y organismos empresariales",
    "Tecnología e innovación": "facilitar pruebas piloto, validación tecnológica y adopción de soluciones con evidencia de desempeño",
    "Mercado y demanda": "fortalecer señales de mercado, demanda y condiciones comerciales para hacer viables las soluciones circulares",
    "Proveedores y cadena de valor": "trabajar con proveedores y clientes para cerrar ciclos y reducir barreras a lo largo de la cadena de valor",
    "Información y medición": "mejorar la disponibilidad de datos, métricas y evidencia para respaldar decisiones e inversiones",
    "Infraestructura": "evaluar infraestructura compartida o especializada que permita escalar soluciones que una empresa no puede resolver por sí sola",
    "Materiales y residuos": "priorizar flujos de materiales y residuos con potencial de valorización, intercambio o reducción",
    "Agua y energía": "identificar proyectos de eficiencia y circularidad en agua y energía con potencial de implementación y medición",
}


def _clean_text(value):
    if value is None:
        return ""
    s = str(value).strip()
    if s.lower() in {"nan", "none", "nat"}:
        return ""
    return re.sub(r"\s+", " ", s)


def _split_multi(value):
    s = _clean_text(value)
    if not s:
        return []
    parts = re.split(r"\s*\|\s*|\s*;\s*", s)
    return [p.strip() for p in parts if p.strip()]


def _top_values(df, column, n=5):
    if df.empty or column not in df.columns:
        return []
    c = Counter()
    for v in df[column].tolist():
        for part in _split_multi(v):
            if part.lower() not in {"sin clasificar", "no se sabe", "ninguno", "ninguna"}:
                c[part] += 1
    return c.most_common(n)


def _row_text(row):
    fields = [
        "hallazgo", "barrera", "barrera_otra", "ejemplo", "actor", "actor_otro",
        "apoyo_solucion", "frase_clave", "notas", "tipo_hallazgo_gremial",
        "afectacion_gremial", "instrumento_gremial"
    ]
    return " ".join(_clean_text(row.get(c, "")) for c in fields).lower()


def _theme_stats(df):
    stats = []
    if df.empty:
        return stats
    for theme, keywords in THEME_KEYWORDS.items():
        rows = []
        for idx, row in df.iterrows():
            txt = _row_text(row)
            if any(k.lower() in txt for k in keywords):
                rows.append(idx)
        if not rows:
            continue
        sub = df.loc[rows]
        groups = int(sub["grupo"].dropna().astype(str).nunique()) if "grupo" in sub.columns else 0
        mesas = int(sub["mesa"].dropna().astype(str).nunique()) if "mesa" in sub.columns else 0
        stats.append({"tema": theme, "menciones": len(rows), "grupos": groups, "mesas": mesas})
    return sorted(stats, key=lambda x: (x["menciones"], x["grupos"], x["mesas"]), reverse=True)


def _join_items(items):
    items = [str(x) for x in items if x]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} y {items[1]}"
    return ", ".join(items[:-1]) + f" y {items[-1]}"


def _scope_signature(df, scope):
    if df.empty:
        raw = scope + "|empty"
    else:
        ids = df["id"].astype(str).tolist() if "id" in df.columns else [str(i) for i in df.index]
        raw = scope + "|" + "|".join(ids)
        if "created_at" in df.columns:
            raw += "|" + "|".join(df["created_at"].fillna("").astype(str).tolist())
        elif "fecha_hora" in df.columns:
            raw += "|" + "|".join(df["fecha_hora"].fillna("").astype(str).tolist())
    return hashlib.md5(raw.encode("utf-8")).hexdigest()[:10]


def _shorten(text, limit=220):
    text = _clean_text(text)
    if not text:
        return ""
    text = " ".join(text.split())
    return text if len(text) <= limit else text[:limit - 1].rstrip() + "…"


def _unique_texts(df, columns, limit=5, max_chars=220):
    out = []
    seen = set()
    if df.empty:
        return out
    for _, row in df.iterrows():
        for col in columns:
            if col not in df.columns:
                continue
            val = _shorten(row.get(col, ""), max_chars)
            key = val.lower().strip()
            if val and key not in seen and key not in {"ninguno", "ninguna", "no aplica", "n/a"}:
                seen.add(key)
                out.append(val)
                if len(out) >= limit:
                    return out
    return out


def _sector_views(df):
    views = []
    if df.empty or "grupo" not in df.columns:
        return views
    groups = [g for g in df["grupo"].dropna().astype(str).unique().tolist() if g.strip()]
    for group in groups:
        gd = df[df["grupo"].astype(str) == group]
        bullets = _unique_texts(
            gd,
            ["hallazgo", "tipo_hallazgo_gremial", "ejemplo", "instrumento_gremial"],
            limit=4,
            max_chars=235,
        )
        barriers = _top_values(gd, "barrera", 2)
        supports = _top_values(gd, "apoyo_solucion", 2)
        if barriers:
            bullets.append("Barreras señaladas: " + _join_items([x[0] for x in barriers]) + ".")
        if supports:
            bullets.append("Vías de avance mencionadas: " + _join_items([x[0] for x in supports]) + ".")
        views.append((group, bullets[:6]))
    return views


def _cross_sector_coincidences(df, themes):
    out = []
    if df.empty:
        return out
    group_total = int(df["grupo"].dropna().astype(str).nunique()) if "grupo" in df.columns else 0
    for t in themes:
        if t["grupos"] >= 2:
            out.append(
                f"{t['tema']}: aparece en {t['grupos']} sectores/grupos"
                + (f" de {group_total}" if group_total else "") + "."
            )
        if len(out) >= 5:
            break
    return out


def _sector_differences(df):
    out = []
    if df.empty or "grupo" not in df.columns:
        return out
    groups = [g for g in df["grupo"].dropna().astype(str).unique().tolist() if g.strip()]
    for group in groups:
        gd = df[df["grupo"].astype(str) == group]
        others = df[df["grupo"].astype(str) != group]
        gthemes = _theme_stats(gd)
        othemes = {x["tema"] for x in _theme_stats(others)}
        unique = [x for x in gthemes if x["tema"] not in othemes]
        if unique:
            out.append(f"{group}: pone un énfasis particular en {unique[0]['tema'].lower()}.")
        else:
            examples = _unique_texts(gd, ["hallazgo", "ejemplo", "tipo_hallazgo_gremial"], limit=1, max_chars=180)
            if examples:
                out.append(f"{group}: {examples[0]}")
        if len(out) >= 5:
            break
    return out


def _opportunities(df):
    vals = _unique_texts(
        df,
        ["apoyo_solucion", "instrumento_gremial", "actor_otro", "notas"],
        limit=7,
        max_chars=240,
    )
    return vals


def _implications(themes, barriers, supports):
    out = []
    for t in themes[:4]:
        action = ACTION_BY_THEME.get(t["tema"])
        if action:
            out.append(f"{t['tema']}: {action}.")
    if not out and supports:
        out = [f"Profundizar en {x[0].lower()} como posible línea de intervención." for x in supports[:4]]
    if barriers and len(out) < 5:
        out.append("Diseñar el seguimiento priorizando las barreras con mayor recurrencia: " + _join_items([x[0].lower() for x in barriers[:3]]) + ".")
    return out[:5]


def _executive_bullets(df, themes, barriers, supports, sector_views):
    bullets = []
    if themes:
        top = themes[:3]
        for t in top:
            reach = f"{t['grupos']} sector{'es' if t['grupos'] != 1 else ''}"
            bullets.append(f"{t['tema']}: concentra {t['menciones']} menciones y aparece en {reach}.")
    if barriers:
        bullets.append("Las principales fricciones se concentran en " + _join_items([x[0].lower() for x in barriers[:3]]) + ".")
    opp = _unique_texts(df, ["apoyo_solucion", "instrumento_gremial"], limit=1, max_chars=230)
    if opp:
        bullets.append("Una vía de avance planteada de forma concreta es: " + opp[0])
    if sector_views and len(sector_views) > 1:
        bullets.append(f"La mesa recoge perspectivas diferenciadas de {len(sector_views)} sectores/grupos; el análisis separa sus posiciones para evitar presentar un consenso artificial.")
    return bullets[:6]


def analizar_sesion(df, scope_label="Sesión completa"):
    data = df.copy()
    total = len(data)
    groups = int(data["grupo"].dropna().astype(str).nunique()) if not data.empty and "grupo" in data.columns else 0
    mesas = int(data["mesa"].dropna().astype(str).nunique()) if not data.empty and "mesa" in data.columns else 0
    high = int((data["prioridad"].fillna("").astype(str) == "Alta").sum()) if not data.empty and "prioridad" in data.columns else 0
    sectorial = int((data["sectorialidad"].fillna("").astype(str) == "Sí, parece sectorial").sum()) if not data.empty and "sectorialidad" in data.columns else 0

    themes = _theme_stats(data)
    barriers = _top_values(data, "barrera", 6)
    supports = _top_values(data, "apoyo_solucion", 6)
    actors = _top_values(data, "actor", 6)
    sector_views = _sector_views(data)
    coincidences = _cross_sector_coincidences(data, themes)
    differences = _sector_differences(data)
    opportunities = _opportunities(data)
    implications = _implications(themes, barriers, supports)
    executive_bullets = _executive_bullets(data, themes, barriers, supports, sector_views)

    priorities = implications[:]

    mesa_summaries = []
    if not data.empty and "mesa" in data.columns:
        for m in MESAS:
            md = data[data["mesa"] == m]
            if not md.empty:
                m_themes = _theme_stats(md)
                m_barriers = _top_values(md, "barrera", 3)
                m_supports = _top_values(md, "apoyo_solucion", 2)
                m_views = _sector_views(md)
                mesa_summaries.append((m, _executive_bullets(md, m_themes, m_barriers, m_supports, m_views)))

    return {
        "scope": scope_label,
        "total": total,
        "groups": groups,
        "mesas": mesas,
        "high": high,
        "sectorial": sectorial,
        "themes": themes,
        "barriers": barriers,
        "supports": supports,
        "actors": actors,
        "executive_bullets": executive_bullets,
        "executive": "\n".join(f"• {x}" for x in executive_bullets),
        "sector_views": sector_views,
        "coincidences": coincidences,
        "differences": differences,
        "opportunities": opportunities,
        "implications": implications,
        "priorities": priorities,
        "mesa_summaries": mesa_summaries,
    }

def _find_montserrat():
    candidates = [
        ASSETS_DIR / "Montserrat-Regular.ttf",
        ASSETS_DIR / "Montserrat.ttf",
        Path("/usr/share/fonts/truetype/montserrat/Montserrat-Regular.ttf"),
        Path("/usr/local/share/fonts/Montserrat-Regular.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return p
    return None


def _pil_font(size=28, bold=False):
    if ImageFont is None:
        return None
    candidates = []
    if bold:
        candidates += [ASSETS_DIR / "Montserrat-SemiBold.ttf", ASSETS_DIR / "Montserrat-Bold.ttf"]
    candidates += [ASSETS_DIR / "Montserrat-Regular.ttf", Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")]
    for p in candidates:
        if p.exists():
            try:
                return ImageFont.truetype(str(p), size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def crear_grafica_barras(items, titulo, max_items=6):
    if PILImage is None or not items:
        return None
    vals = [(str(k), int(v)) for k, v in items[:max_items] if int(v) > 0]
    if not vals:
        return None
    W = 1500
    left = 420
    right = 120
    top = 140
    row_h = 86
    H = top + row_h * len(vals) + 80
    img = PILImage.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    font_title = _pil_font(36, bold=True)
    font_label = _pil_font(25)
    font_num = _pil_font(24, bold=True)
    draw.text((40, 35), titulo, fill=(23, 59, 46), font=font_title)
    maxv = max(v for _, v in vals)
    for i, (label, value) in enumerate(vals):
        y = top + i * row_h
        wrapped = textwrap.wrap(label, width=33)[:2]
        draw.multiline_text((40, y + 5), "\n".join(wrapped), fill=(45, 60, 51), font=font_label, spacing=3)
        x0 = left
        x1 = left + int((W - left - right) * (value / maxv))
        draw.rounded_rectangle((x0, y + 10, x1, y + 54), radius=18, fill=(47, 110, 74))
        draw.text((x1 + 14, y + 17), str(value), fill=(23, 59, 46), font=font_num)
    bio = io.BytesIO()
    img.save(bio, format="PNG")
    return bio.getvalue()


def _set_docx_montserrat(doc):
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Montserrat"
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
                rfonts.set(qn(f"w:{attr}"), "Montserrat")
        except Exception:
            pass


def _add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    run2 = paragraph.add_run(" de ")
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES")
    run2._r.addnext(fld2)


def _docx_add_logo_row(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.3)
    table.columns[1].width = Inches(3.3)
    c1, c2 = table.rows[0].cells
    c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p1 = c1.paragraphs[0]
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_COINVIERTE and Path(LOGO_COINVIERTE).exists():
        p1.add_run().add_picture(str(LOGO_COINVIERTE), width=Inches(1.85))
    else:
        p1.add_run("COINVIERTE").bold = True
    if LOGO_TEC and Path(LOGO_TEC).exists():
        p2.add_run().add_picture(str(LOGO_TEC), width=Inches(1.3))
    else:
        p2.add_run("Tecnológico de Monterrey").bold = True
    for cell in [c1, c2]:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "nil")
            borders.append(tag)
        tcPr.append(borders)


def generar_docx_sintesis(analysis, executive_text, priorities_text, chart_theme=None, chart_barrier=None):
    if Document is None:
        return None
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    _set_docx_montserrat(doc)
    doc.styles["Normal"].font.size = Pt(9.5)
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(11.5)

    _docx_add_logo_row(doc)
    p = doc.add_paragraph(); p.style = doc.styles["Title"]
    r = p.add_run("Síntesis de sesión"); r.font.color.rgb = RGBColor(23, 59, 46)
    p2 = doc.add_paragraph("Diagnóstico de Economía Circular para el Estado de Jalisco")
    p2.runs[0].bold = True; p2.runs[0].font.size = Pt(11); p2.runs[0].font.color.rgb = RGBColor(47, 110, 74)
    mesa_txt = analysis['scope'] if analysis['scope'] != 'Sesión completa' else 'Sesión completa'
    meta = doc.add_paragraph(f"Mesa: {mesa_txt}  |  Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    meta.runs[0].font.size = Pt(8.5); meta.runs[0].font.color.rgb = RGBColor(110, 123, 115)

    def bullets(title, items):
        if not items: return
        doc.add_heading(title, level=1)
        for item in items:
            item = str(item).strip().lstrip('•- ').strip()
            if item: doc.add_paragraph(item, style="List Bullet")

    bullets("Lectura ejecutiva", [x for x in executive_text.splitlines() if x.strip()])

    if analysis.get("sector_views"):
        doc.add_heading("Perspectiva por sector", level=1)
        for sector, items in analysis["sector_views"]:
            doc.add_heading(sector, level=2)
            for item in items:
                doc.add_paragraph(item, style="List Bullet")

    bullets("Coincidencias entre sectores", analysis.get("coincidences", []))
    bullets("Diferencias o énfasis sectoriales", analysis.get("differences", []))
    bullets("Oportunidades identificadas", analysis.get("opportunities", []))
    bullets("Implicaciones para COINVIERTE", [x for x in priorities_text.splitlines() if x.strip()])

    if chart_theme or chart_barrier:
        doc.add_heading("Gráficas de apoyo", level=1)
        if chart_theme: doc.add_picture(io.BytesIO(chart_theme), width=Inches(6.7))
        if chart_barrier: doc.add_picture(io.BytesIO(chart_barrier), width=Inches(6.7))

    doc.add_heading("Nota metodológica", level=1)
    doc.add_paragraph("Esta síntesis interpreta y agrupa los registros capturados por mesa y sector. No atribuye comentarios a personas individuales. Las coincidencias y diferencias se construyen a partir de los registros disponibles; las frecuencias representan menciones y no proporciones de participantes.")

    footer = sec.footer
    fp = footer.paragraphs[0]; fp.text = "COINVIERTE · Diagnóstico de Economía Circular Jalisco"
    fp.runs[0].font.name = "Montserrat"; fp.runs[0].font.size = Pt(7.5)
    _add_page_field(footer.add_paragraph())
    out = io.BytesIO(); doc.save(out); return out.getvalue()


class _NumberedCanvas(rl_canvas.Canvas if colors is not None else object):
    def __init__(self, *args, **kwargs):
        if colors is None:
            return
        rl_canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_page_number(page_count)
            rl_canvas.Canvas.showPage(self)
        rl_canvas.Canvas.save(self)

    def _draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 7.5)
        self.setFillColorRGB(0.43, 0.48, 0.45)
        self.drawRightString(letter[0] - 17 * mm, 10 * mm, f"Página {self._pageNumber} de {page_count}")
        self.drawString(17 * mm, 10 * mm, "COINVIERTE · Diagnóstico de Economía Circular Jalisco")
        self.restoreState()


def _register_pdf_font():
    regular = _find_montserrat()
    if regular is not None and colors is not None:
        try:
            pdfmetrics.registerFont(TTFont("Montserrat", str(regular)))
            return "Montserrat"
        except Exception:
            pass
    return "Helvetica"


def generar_pdf_sintesis(analysis, executive_text, priorities_text, chart_theme=None, chart_barrier=None):
    if colors is None:
        return None
    out = io.BytesIO()
    font = _register_pdf_font()
    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleEC", parent=styles["Title"], fontName=font, fontSize=19, leading=23, textColor=colors.HexColor("#173B2E"), spaceAfter=4)
    subtitle = ParagraphStyle("SubtitleEC", parent=styles["Normal"], fontName=font, fontSize=10.5, leading=14, textColor=colors.HexColor("#2F6E4A"), spaceAfter=10)
    h1 = ParagraphStyle("H1EC", parent=styles["Heading1"], fontName=font, fontSize=13, leading=16, textColor=colors.HexColor("#173B2E"), spaceBefore=9, spaceAfter=5)
    h2 = ParagraphStyle("H2EC", parent=styles["Heading2"], fontName=font, fontSize=10.5, leading=13, textColor=colors.HexColor("#285843"), spaceBefore=7, spaceAfter=3)
    body = ParagraphStyle("BodyEC", parent=styles["BodyText"], fontName=font, fontSize=9.2, leading=13.2, textColor=colors.HexColor("#2D3C33"), spaceAfter=5)
    bullet = ParagraphStyle("BulletEC", parent=body, leftIndent=12, firstLineIndent=-7, bulletIndent=3, spaceAfter=3)
    meta = ParagraphStyle("MetaEC", parent=body, fontSize=8, textColor=colors.HexColor("#6E7B73"), spaceAfter=8)

    def esc(x):
        return str(x).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def add_bullets(story, heading, items):
        vals = [str(x).strip().lstrip('•- ').strip() for x in items if str(x).strip()]
        if not vals:
            return
        story.append(Paragraph(esc(heading), h1))
        for item in vals:
            story.append(Paragraph("• " + esc(item), bullet))

    doc = SimpleDocTemplate(out, pagesize=letter, rightMargin=17*mm, leftMargin=17*mm, topMargin=16*mm, bottomMargin=18*mm, title="Síntesis de sesión - Economía Circular Jalisco")
    story = []
    logo_cells = []
    if LOGO_COINVIERTE and Path(LOGO_COINVIERTE).exists():
        logo_cells.append(RLImage(str(LOGO_COINVIERTE), width=42*mm, height=15*mm, kind="proportional"))
    else:
        logo_cells.append(Paragraph("<b>COINVIERTE</b>", body))
    if LOGO_TEC and Path(LOGO_TEC).exists():
        logo_cells.append(RLImage(str(LOGO_TEC), width=31*mm, height=13*mm, kind="proportional"))
    else:
        logo_cells.append(Paragraph("<b>Tecnológico de Monterrey</b>", body))
    lt = Table([logo_cells], colWidths=[85*mm, 85*mm])
    lt.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("ALIGN", (1,0), (1,0), "RIGHT"), ("BOTTOMPADDING", (0,0), (-1,-1), 6)]))
    mesa_txt = analysis['scope'] if analysis['scope'] != 'Sesión completa' else 'Sesión completa'
    story += [lt, Paragraph("Síntesis de sesión", title), Paragraph("Diagnóstico de Economía Circular para el Estado de Jalisco", subtitle), Paragraph(f"Mesa: {esc(mesa_txt)} &nbsp;&nbsp;|&nbsp;&nbsp; Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", meta)]

    add_bullets(story, "Lectura ejecutiva", executive_text.splitlines())

    if analysis.get("sector_views"):
        story.append(Paragraph("Perspectiva por sector", h1))
        for sector, items in analysis["sector_views"]:
            story.append(Paragraph(esc(sector), h2))
            for item in items:
                story.append(Paragraph("• " + esc(item), bullet))

    add_bullets(story, "Coincidencias entre sectores", analysis.get("coincidences", []))
    add_bullets(story, "Diferencias o énfasis sectoriales", analysis.get("differences", []))
    add_bullets(story, "Oportunidades identificadas", analysis.get("opportunities", []))
    add_bullets(story, "Implicaciones para COINVIERTE", priorities_text.splitlines())

    if chart_theme or chart_barrier:
        story.append(Paragraph("Gráficas de apoyo", h1))
        if chart_theme:
            story += [RLImage(io.BytesIO(chart_theme), width=175*mm, height=70*mm, kind="proportional"), Spacer(1, 4*mm)]
        if chart_barrier:
            story += [RLImage(io.BytesIO(chart_barrier), width=175*mm, height=70*mm, kind="proportional"), Spacer(1, 4*mm)]

    story += [Paragraph("Nota metodológica", h1), Paragraph("Esta síntesis interpreta y agrupa los registros capturados por mesa y sector. No atribuye comentarios a personas individuales. Las coincidencias y diferencias se construyen a partir de los registros disponibles; las frecuencias representan menciones y no proporciones de participantes.", body)]
    doc.build(story, canvasmaker=_NumberedCanvas)
    return out.getvalue()

def info_card(icon, icon_class, label, value):
    st.markdown(
        f"""
        <div class="info-card">
            <div class="info-icon {icon_class}">{icon}</div>
            <div>
                <div class="info-label">{label}</div>
                <div class="info-value">{value}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ==========================================================
# Estado compartido entre Hallazgos y Resumen
# ==========================================================
if "mesa_contexto" not in st.session_state or st.session_state["mesa_contexto"] not in MESAS:
    st.session_state["mesa_contexto"] = MESAS[0]
if "mesa_resumen" not in st.session_state or st.session_state["mesa_resumen"] not in MESAS:
    st.session_state["mesa_resumen"] = st.session_state["mesa_contexto"]

def sync_mesa_from_sidebar():
    mesa_sel = st.session_state.get("mesa_contexto")
    if mesa_sel in MESAS:
        st.session_state["mesa_resumen"] = mesa_sel

def sync_mesa_from_hallazgos():
    mesa_sel = st.session_state.get("f_mesa")
    if mesa_sel in MESAS:
        st.session_state["mesa_resumen"] = mesa_sel

def nuevo_hallazgo():
    for key, default in {
        "barreras_live": [],
        "barrera_otra_live": "",
        "actores_live": [],
        "actor_otro_live": "",
        "tipo_hallazgo_gremial_live": [],
        "afectacion_gremial_live": [],
        "actores_gremiales_live": [],
        "instrumentos_gremiales_live": [],
        "otro_tipo_gremial_live": "",
        "otro_afectado_gremial_live": "",
        "otro_actor_gremial_live": "",
        "otro_instrumento_gremial_live": "",
    }.items():
        st.session_state[key] = default
    st.session_state["hallazgo_guardado"] = False


# ==========================================================
# Sidebar
# ==========================================================
st.sidebar.markdown(
    """
    <div style="margin-bottom:1rem;">
        <div style="font-size:1.15rem;font-weight:850;color:#173B2E;">Relatoría</div>
        <div style="font-size:.96rem;font-weight:700;color:#55675D;">Economía Circular Jalisco</div>
    </div>
    """,
    unsafe_allow_html=True,
)
st.sidebar.markdown("## Datos de la ronda")

mesa = st.sidebar.selectbox(
    "Mesa temática",
    MESAS,
    key="mesa_contexto",
    on_change=sync_mesa_from_sidebar,
)
st.sidebar.caption(f"Mesa seleccionada: {mesa}")
grupo = st.sidebar.selectbox("Grupo / sector", GRUPOS)
ronda = st.sidebar.selectbox("Ronda", [1, 2, 3, 4])
relator = st.sidebar.text_input("Relator/a", placeholder="Nombre")
moderador = st.sidebar.text_input("Moderador/a", placeholder="Nombre")

st.sidebar.divider()

st.sidebar.markdown(
    """
    <div class="sidebar-note">
        👥 La mesa, moderador y relator permanecen anclados.
        Los grupos rotan.
    </div>
    """,
    unsafe_allow_html=True,
)

if grupo == "Líderes gremiales":
    st.sidebar.info(
        "Modo gremial activo: preguntas con enfoque sectorial y de política pública."
    )

if supabase:
    st.sidebar.markdown(
        """
        <div class="sidebar-db">
            🗄️ Base compartida:<br>
            <span style="font-size:1.05rem;">Supabase</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
else:
    st.sidebar.warning("Modo local")


# ==========================================================
# Encabezado
# ==========================================================

# Franja superior con las dos instituciones
with st.container(border=True):
    b1, bd, b2, spacer = st.columns(
        [1.15, 0.08, 0.95, 3.2],
        vertical_alignment="center"
    )

    with b1:
        if LOGO_COINVIERTE is not None:
            st.image(str(LOGO_COINVIERTE), width=190)
        else:
            st.markdown(
                f"<div style='color:{MUTED};font-weight:800;'>COINVIERTE</div>",
                unsafe_allow_html=True,
            )

    with bd:
        st.markdown(
            '<div class="brand-separator"></div>',
            unsafe_allow_html=True,
        )

    with b2:
        if LOGO_TEC is not None:
            # Mantener proporciones del archivo oficial sin alterar/redibujar el logotipo.
            st.image(str(LOGO_TEC), width=120)
        else:
            st.markdown(
                f"<div style='color:{MUTED};font-weight:800;'>Tecnológico de Monterrey</div>",
                unsafe_allow_html=True,
            )

    with spacer:
        st.write("")

st.write("")

# Hero principal
hero_html = f"""<div class="hero-shell">
<div style="display:inline-block;background:{SAGE_2};color:{GREEN_DARK};border-radius:999px;padding:.38rem .75rem;font-size:.78rem;font-weight:800;margin-bottom:.8rem;">🌿 Relatoría</div>
<div class="hero-title">Economía Circular Jalisco</div>
<div style="width:54px;height:4px;border-radius:99px;background:{GREEN};margin:.7rem 0 1rem 0;"></div>
<div class="hero-sub">Captura ágil de hallazgos para identificar barreras, oportunidades, actores habilitadores y soluciones con potencial de escalamiento.</div>
<div class="pill-row">
<span class="soft-pill pill-sage">🌿 {mesa}</span>
<span class="soft-pill pill-sand">👥 {grupo}</span>
<span class="soft-pill pill-aqua">🗓️ Ronda {ronda}</span>
</div>
</div>"""

st.markdown(hero_html, unsafe_allow_html=True)

st.write("")

# Context cards compactas
c1, c2, c3, c4 = st.columns(4)

with c1:
    info_card("👥", "icon-sage", "Mesa activa", mesa)

with c2:
    info_card("👤", "icon-sand", "Grupo actual", grupo)

with c3:
    info_card("🗓️", "icon-aqua", "Ronda", str(ronda))

with c4:
    info_card(
        "🗄️",
        "icon-sage",
        "Base de datos",
        "Supabase" if supabase else "Local",
    )


# ==========================================================
# Tabs
# ==========================================================
tab_captura, tab_hallazgos, tab_resumen = st.tabs(
    ["✎ Captura", "☷ Hallazgos", "▥ Síntesis"]
)


# ==========================================================
# Captura
# ==========================================================
with tab_captura:
    st.markdown("## Registrar hallazgo")

    st.markdown(
        """
        <div class="rule-note">
            💡 <b>Regla simple:</b> un registro = un hallazgo.
            Captura el punto útil, la barrera y qué podría destrabarlo.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Referencia de la conversación")
    preguntas_activas = (
        PREGUNTAS_LIDERES_GREMIALES[mesa]
        if grupo == "Líderes gremiales"
        else PREGUNTAS_POR_MESA[mesa]
    )

    pregunta_referencia = st.selectbox(
        "Pregunta de referencia *",
        preguntas_activas,
        help=(
            "La lista cambia según la mesa y el tipo de grupo. "
            "Para líderes gremiales se usa una batería sectorial y de política pública."
        )
    )

    if grupo == "Líderes gremiales":
        st.caption(
            "Bloque para líderes gremiales: enfocado en patrones sectoriales, barreras sistémicas, "
            "instrumentos, coordinación y condiciones para escalar la economía circular."
        )
    else:
        st.caption(
            "Selecciona la pregunta que originó el hallazgo. "
            "Así podremos sistematizar después por mesa, pregunta, sector y barrera."
        )

    es_cierre = pregunta_referencia.startswith("Cierre.")

    if es_cierre:
        # Cierre: una síntesis directa, sin clasificaciones adicionales.
        with st.form("cierre_form", clear_on_submit=True):
            respuesta_cierre = st.text_area(
                "Respuesta / síntesis de cierre *",
                placeholder="Captura aquí la conclusión principal del grupo para esta pregunta de cierre.",
                height=150,
            )

            guardar_cierre = st.form_submit_button(
                "💾 Guardar cierre",
                type="primary",
                use_container_width=True,
            )

            if guardar_cierre:
                if not respuesta_cierre.strip():
                    st.error("Captura la respuesta o síntesis de cierre.")
                elif not relator.strip():
                    st.error("Escribe el nombre del relator/a en la barra lateral.")
                else:
                    rec = {
                        "fecha_hora": datetime.now().isoformat(timespec="seconds"),
                        "mesa": mesa,
                        "grupo": grupo,
                        "ronda": int(ronda),
                        "relator": relator.strip(),
                        "moderador": moderador.strip(),
                        "tipo_pregunta": "",
                        "pregunta_referencia": pregunta_referencia,
                        "hallazgo": respuesta_cierre.strip(),
                        "barrera": "",
                        "barrera_otra": "",
                        "ejemplo": "",
                        "actor": "",
                        "actor_otro": "",
                        "apoyo_solucion": "",
                        "sectorialidad": "",
                        "prioridad": "",
                        "frase_clave": "",
                        "notas": "",
                        "tipo_hallazgo_gremial": "",
                        "afectacion_gremial": "",
                        "instrumento_gremial": "",
                    }

                    ok, msg = save_record(rec)
                    if ok:
                        _msg_guardado = st.success("Cierre guardado correctamente.")
                    else:
                        _msg_error = st.error(msg)

    elif grupo == "Líderes gremiales":
        # Formulario especial: patrón sectorial + política pública.
        st.markdown(
            """
            <div style="
                background:#EEF5E8;
                border:1px solid #DDE5DA;
                border-radius:12px;
                padding:10px 13px;
                margin-bottom:12px;">
                <b>Captura gremial:</b> registra el patrón sectorial, a quién afecta,
                qué actores deben intervenir y qué instrumento podría destrabarlo.
            </div>
            """,
            unsafe_allow_html=True,
        )

        tipo_hallazgo_sel = st.multiselect(
            "Tipo de hallazgo",
            TIPOS_HALLAZGO_GREMIAL,
            placeholder="Selecciona una o varias categorías",
            key="tipo_hallazgo_gremial_live",
        )

        afectacion_sel = st.multiselect(
            "¿A quién afecta principalmente?",
            AFECTACION_GREMIAL,
            placeholder="Selecciona una o varias opciones",
            key="afectacion_gremial_live",
        )

        actores_gremiales_sel = st.multiselect(
            "Actores que tendrían que intervenir",
            ACTORES_GREMIALES,
            placeholder="Selecciona uno o varios actores",
            key="actores_gremiales_live",
        )

        instrumentos_sel = st.multiselect(
            "Instrumento o acción que podría destrabarlo",
            INSTRUMENTOS_GREMIALES,
            placeholder="Selecciona una o varias opciones",
            key="instrumentos_gremiales_live",
        )

        otro_tipo = st.text_input(
            "Otro tipo de hallazgo",
            placeholder="Agrega otra categoría o precisión",
            key="otro_tipo_gremial_live",
        )

        otro_afectado = st.text_input(
            "Otro grupo afectado",
            placeholder="Agrega otro grupo, subsector o precisión",
            key="otro_afectado_gremial_live",
        )

        otro_actor_gremial = st.text_input(
            "Otro actor",
            placeholder="Agrega otro actor o precisión",
            key="otro_actor_gremial_live",
        )

        otro_instrumento = st.text_input(
            "Otro instrumento o acción",
            placeholder="Agrega otro instrumento, acción o precisión",
            key="otro_instrumento_gremial_live",
        )

        with st.form("hallazgo_gremial_form", clear_on_submit=True):
            hallazgo_gremial = st.text_area(
                "Hallazgo / patrón sectorial *",
                placeholder="¿Qué patrón, problema u oportunidad identifica en el sector?",
                height=105,
            )

            evidencia_gremial = st.text_area(
                "Ejemplo o evidencia sectorial",
                placeholder="Caso recurrente, dato, experiencia de afiliados o diferencia entre tipos de empresa.",
                height=90,
            )

            prioridad_gremial = st.selectbox(
                "Prioridad",
                ["Sin clasificar", "Alta", "Media", "Baja"],
            )

            guardar_gremial = st.form_submit_button(
                "💾 Guardar hallazgo gremial",
                type="primary",
                use_container_width=True,
            )

            if guardar_gremial:
                if not hallazgo_gremial.strip():
                    st.error("Captura el hallazgo o patrón sectorial antes de guardar.")
                elif not relator.strip():
                    st.error("Escribe el nombre del relator/a en la barra lateral.")
                else:
                    tipo_txt = " | ".join(tipo_hallazgo_sel)
                    if otro_tipo.strip():
                        tipo_txt = f"{tipo_txt} | Otro: {otro_tipo.strip()}" if tipo_txt else f"Otro: {otro_tipo.strip()}"

                    afectacion_txt = " | ".join(afectacion_sel)
                    if otro_afectado.strip():
                        afectacion_txt = f"{afectacion_txt} | Otro: {otro_afectado.strip()}" if afectacion_txt else f"Otro: {otro_afectado.strip()}"

                    actor_txt = " | ".join(actores_gremiales_sel)
                    if otro_actor_gremial.strip():
                        actor_txt = f"{actor_txt} | Otro: {otro_actor_gremial.strip()}" if actor_txt else f"Otro: {otro_actor_gremial.strip()}"

                    instrumento_txt = " | ".join(instrumentos_sel)
                    if otro_instrumento.strip():
                        instrumento_txt = f"{instrumento_txt} | Otro: {otro_instrumento.strip()}" if instrumento_txt else f"Otro: {otro_instrumento.strip()}"

                    rec = {
                        "fecha_hora": datetime.now().isoformat(timespec="seconds"),
                        "mesa": mesa,
                        "grupo": grupo,
                        "ronda": int(ronda),
                        "relator": relator.strip(),
                        "moderador": moderador.strip(),
                        "tipo_pregunta": "",
                        "pregunta_referencia": pregunta_referencia,
                        "hallazgo": hallazgo_gremial.strip(),
                        "barrera": tipo_txt,
                        "barrera_otra": "",
                        "ejemplo": evidencia_gremial.strip(),
                        "actor": actor_txt,
                        "actor_otro": "",
                        "apoyo_solucion": instrumento_txt,
                        "sectorialidad": afectacion_txt,
                        "prioridad": prioridad_gremial,
                        "frase_clave": "",
                        "notas": "",
                        "tipo_hallazgo_gremial": tipo_txt,
                        "afectacion_gremial": afectacion_txt,
                        "instrumento_gremial": instrumento_txt,
                    }

                    ok, msg = save_record(rec)
                    if ok:
                        st.session_state["hallazgo_guardado"] = True
                    else:
                        _msg_error = st.error(msg)

        if st.session_state.get("hallazgo_guardado"):
            st.success("Hallazgo gremial guardado correctamente.")
            st.button(
                "+ Registrar nuevo hallazgo",
                type="primary",
                use_container_width=True,
                on_click=nuevo_hallazgo,
                key="btn_nuevo_hallazgo",
            )

    else:
        # Formulario estándar para sectores primario, secundario y terciario.
        pre1, pre2 = st.columns(2)

        with pre1:
            barreras_sel = st.multiselect(
                "Barreras",
                BARRERAS[1:],
                placeholder="Selecciona una o varias barreras",
                key="barreras_live"
            )
            barrera_otra = st.text_input(
                "Otra barrera",
                placeholder="Agrega otra barrera o precisión, aunque ya hayas seleccionado una categoría",
                key="barrera_otra_live"
            )

        with pre2:
            actores_sel = st.multiselect(
                "Actores que pueden habilitar",
                ACTORES[1:],
                placeholder="Selecciona uno o varios actores",
                key="actores_live"
            )
            actor_otro = st.text_input(
                "Otro actor",
                placeholder="Agrega otro actor o precisión, aunque ya hayas seleccionado uno",
                key="actor_otro_live"
            )

        with st.form("hallazgo_form", clear_on_submit=True):
            hallazgo = st.text_area(
                "Problema / oportunidad *",
                placeholder=(
                    "Ej. Se generan subproductos orgánicos sin una salida "
                    "comercial estable."
                ),
                height=100,
            )

            c1, c2 = st.columns(2)

            with c1:
                prioridad = st.selectbox(
                    "Prioridad percibida",
                    PRIORIDADES
                )

            with c2:
                sectorialidad = st.selectbox(
                    "¿Parece sectorial?",
                    SECTORIALIDAD
                )

            e1, e2 = st.columns(2)

            with e1:
                ejemplo = st.text_area(
                    "Ejemplo concreto",
                    placeholder=(
                        "Describe un caso, dato o situación que ilustre "
                        "el hallazgo."
                    ),
                    height=88,
                )

            with e2:
                apoyo = st.text_area(
                    "Apoyo / solución sugerida",
                    placeholder=(
                        "¿Qué apoyo, herramienta o solución podría ayudar "
                        "a destrabarlo?"
                    ),
                    height=88,
                )

            with st.expander("Campos adicionales"):
                x1, x2 = st.columns(2)
                with x1:
                    frase = st.text_area(
                        "Frase clave (sin atribución)",
                        height=70
                    )
                with x2:
                    notas = st.text_area(
                        "Notas breves",
                        height=70
                    )

            b1, b2 = st.columns([1, 2])

            with b1:
                limpiar = st.form_submit_button(
                    "↻ Limpiar formulario",
                    use_container_width=True,
                )

            with b2:
                guardar = st.form_submit_button(
                    "💾 Guardar hallazgo",
                    type="primary",
                    use_container_width=True,
                )

            if limpiar:
                st.rerun()

            if guardar:
                if not hallazgo.strip():
                    st.error(
                        "Captura el problema u oportunidad antes de guardar."
                    )

                elif not relator.strip():
                    st.error(
                        "Escribe el nombre del relator/a en la barra lateral."
                    )

                else:
                    barrera_txt = " | ".join(barreras_sel)
                    actor_txt = " | ".join(actores_sel)

                    rec = {
                        "fecha_hora": datetime.now().isoformat(
                            timespec="seconds"
                        ),
                        "mesa": mesa,
                        "grupo": grupo,
                        "ronda": int(ronda),
                        "relator": relator.strip(),
                        "moderador": moderador.strip(),
                        "tipo_pregunta": "",
                        "pregunta_referencia": pregunta_referencia,
                        "hallazgo": hallazgo.strip(),
                        "barrera": barrera_txt,
                        "barrera_otra": barrera_otra.strip(),
                        "ejemplo": ejemplo.strip(),
                        "actor": actor_txt,
                        "actor_otro": actor_otro.strip(),
                        "apoyo_solucion": apoyo.strip(),
                        "sectorialidad": sectorialidad,
                        "prioridad": prioridad,
                        "frase_clave": frase.strip(),
                        "notas": notas.strip(),
                        "tipo_hallazgo_gremial": "",
                        "afectacion_gremial": "",
                        "instrumento_gremial": "",
                    }

                    ok, msg = save_record(rec)

                    if ok:
                        st.session_state["hallazgo_guardado"] = True
                    else:
                        _msg_error = st.error(msg)

        if st.session_state.get("hallazgo_guardado"):
            st.success("Hallazgo guardado correctamente.")
            st.button(
                "+ Registrar nuevo hallazgo",
                type="primary",
                use_container_width=True,
                on_click=nuevo_hallazgo,
                key="btn_nuevo_hallazgo",
            )

    if grupo != "Líderes gremiales":
        st.markdown(
            """
            <div style="
                margin-top:14px;
                background:#FFF7E6;
                border-left:4px solid #D4A33D;
                border-radius:12px;
                padding:11px 14px;
                color:#5E543E;">
                <b>Si una empresa no está haciendo nada:</b>
                también es un hallazgo. Si surge en la conversación,
                registra por qué no ha empezado: prioridad, conocimiento,
                presupuesto, liderazgo, proveedores, información u otra causa.
            </div>
            """,
            unsafe_allow_html=True,
        )


# ==========================================================
# Hallazgos
# ==========================================================
with tab_hallazgos:
    df = load_data()

    st.markdown("## Hallazgos registrados")

    m1, m2, m3, m4 = st.columns(4)

    total = len(df)
    altas = (
        int((df["prioridad"] == "Alta").sum())
        if not df.empty and "prioridad" in df.columns
        else 0
    )
    sectoriales = (
        int(
            (
                df["sectorialidad"]
                == "Sí, parece sectorial"
            ).sum()
        )
        if not df.empty and "sectorialidad" in df.columns
        else 0
    )
    mesas_con_datos = (
        int(df["mesa"].nunique())
        if not df.empty and "mesa" in df.columns
        else 0
    )

    m1.metric("Registros", total)
    m2.metric("Prioridad alta", altas)
    m3.metric("Hallazgos sectoriales", sectoriales)
    m4.metric("Mesas con hallazgos", mesas_con_datos)

    f1, f2, f3, f4 = st.columns(4)

    with f1:
        filtro_mesa = st.selectbox(
            "Filtrar por mesa",
            ["Todas"] + MESAS,
            key="f_mesa",
            on_change=sync_mesa_from_hallazgos,
            help="Al elegir una mesa, esa selección se conserva en Resumen de mesa.",
        )

    with f2:
        filtro_grupo = st.selectbox(
            "Filtrar por grupo",
            ["Todos"] + GRUPOS,
            key="f_grupo"
        )

    preguntas_filtro = []
    if filtro_mesa != "Todas":
        preguntas_filtro = (
            PREGUNTAS_LIDERES_GREMIALES[filtro_mesa]
            if filtro_grupo == "Líderes gremiales"
            else PREGUNTAS_POR_MESA[filtro_mesa]
        )
    elif not df.empty and "pregunta_referencia" in df.columns:
        preguntas_filtro = sorted(
            [
                x for x in
                df["pregunta_referencia"].dropna().astype(str).unique().tolist()
                if x
            ]
        )

    with f3:
        filtro_pregunta = st.selectbox(
            "Filtrar por pregunta",
            ["Todas"] + preguntas_filtro,
            key="f_pregunta"
        )

    with f4:
        filtro_barrera = st.selectbox(
            "Filtrar por barrera",
            ["Todas"] + BARRERAS[1:],
            key="f_barrera"
        )

    view = df.copy()

    if not view.empty:
        if filtro_mesa != "Todas":
            view = view[
                view["mesa"] == filtro_mesa
            ]

        if filtro_grupo != "Todos":
            view = view[
                view["grupo"] == filtro_grupo
            ]


        if (
            filtro_pregunta != "Todas"
            and "pregunta_referencia" in view.columns
        ):
            view = view[
                view["pregunta_referencia"] == filtro_pregunta
            ]

        if filtro_barrera != "Todas":
            view = view[
                view["barrera"].fillna("").astype(str).str.contains(
                    re.escape(filtro_barrera),
                    regex=True
                )
            ]

    st.dataframe(
        view,
        use_container_width=True,
        hide_index=True,
        column_config={
            "pregunta_referencia": st.column_config.TextColumn(
                "Pregunta de referencia",
                width="large"
            ),
            "tipo_hallazgo_gremial": st.column_config.TextColumn(
                "Tipo de hallazgo gremial",
                width="large"
            ),
            "afectacion_gremial": st.column_config.TextColumn(
                "Afectación gremial",
                width="medium"
            ),
            "instrumento_gremial": st.column_config.TextColumn(
                "Instrumento / acción gremial",
                width="large"
            ),
            "hallazgo": st.column_config.TextColumn(
                "Hallazgo",
                width="large"
            ),
            "apoyo_solucion": st.column_config.TextColumn(
                "Apoyo / solución",
                width="large"
            ),
        },
    )

    d1, d2 = st.columns(2)

    with d1:
        st.download_button(
            "Descargar CSV",
            data=view.to_csv(
                index=False
            ).encode("utf-8-sig"),
            file_name="relatoria_economia_circular.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with d2:
        st.download_button(
            "Descargar Excel",
            data=to_excel_bytes(view),
            file_name="relatoria_economia_circular.xlsx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
            use_container_width=True,
        )

    with st.expander("Borrar un registro"):
        if view.empty or "id" not in view.columns:
            st.caption(
                "No hay registros disponibles."
            )
        else:
            ids = view["id"].dropna().tolist()
            rid = st.selectbox(
                "ID del registro",
                ids
            )

            if st.button("Borrar registro"):
                if delete_record(rid):
                    _msg_borrado = st.success(
                        "Registro borrado."
                    )
                    st.rerun()


# ==========================================================
# Síntesis
# ==========================================================
with tab_resumen:
    df = load_data()

    st.markdown("## Síntesis de la sesión")
    st.caption(
        "Esta vista no transcribe los registros. Agrupa patrones, recurrencias, barreras y posibles líneas de seguimiento para facilitar una lectura ejecutiva."
    )

    s1, s2 = st.columns([1, 1.7])
    with s1:
        alcance = st.radio(
            "Alcance de la síntesis",
            ["Sesión completa", "Mesa específica"],
            horizontal=True,
            key="alcance_sintesis",
        )
    with s2:
        if alcance == "Mesa específica":
            mesa_resumen = st.selectbox(
                "Mesa a sintetizar",
                MESAS,
                key="mesa_resumen",
                help="Si filtraste una mesa en Hallazgos, esa selección se conserva.",
            )
        else:
            mesa_resumen = None
            st.markdown("<div style='height:1.75rem'></div>", unsafe_allow_html=True)
            st.markdown("**Incluye todas las mesas con registros.**")

    if alcance == "Mesa específica":
        synth_df = (
            df[df["mesa"] == mesa_resumen].copy()
            if not df.empty and "mesa" in df.columns
            else pd.DataFrame()
        )
        scope_label = mesa_resumen
    else:
        synth_df = df.copy()
        scope_label = "Sesión completa"

    analysis = analizar_sesion(synth_df, scope_label=scope_label)

    if synth_df.empty:
        st.info("Todavía no hay registros suficientes para generar la síntesis seleccionada.")
    else:
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Hallazgos analizados", analysis["total"])
        m2.metric("Grupos escuchados", analysis["groups"])
        m3.metric("Prioridad alta", analysis["high"])
        m4.metric("Hallazgos sectoriales", analysis["sectorial"])

        signature = _scope_signature(synth_df, scope_label)
        exec_key = f"sintesis_exec_{signature}"
        pri_key = f"sintesis_pri_{signature}"
        if exec_key not in st.session_state:
            st.session_state[exec_key] = analysis["executive"]
        if pri_key not in st.session_state:
            st.session_state[pri_key] = "\n".join(f"• {x}" for x in analysis["priorities"])

        st.markdown(f"### Mesa: {scope_label}")
        st.markdown("#### Lectura ejecutiva")
        for item in analysis["executive_bullets"]:
            st.markdown(f"- {item}")

        if analysis["sector_views"]:
            st.markdown("#### Perspectiva por sector")
            for sector, items in analysis["sector_views"]:
                st.markdown(f"**{sector}**")
                for item in items:
                    st.markdown(f"- {item}")

        if analysis["coincidences"]:
            st.markdown("#### Coincidencias entre sectores")
            for item in analysis["coincidences"]:
                st.markdown(f"- {item}")

        if analysis["differences"]:
            st.markdown("#### Diferencias o énfasis sectoriales")
            for item in analysis["differences"]:
                st.markdown(f"- {item}")

        if analysis["opportunities"]:
            st.markdown("#### Oportunidades identificadas")
            for item in analysis["opportunities"]:
                st.markdown(f"- {item}")

        st.markdown("#### Implicaciones para COINVIERTE")
        if analysis["implications"]:
            for item in analysis["implications"]:
                st.markdown(f"- {item}")
        else:
            st.caption("Todavía no hay suficientes registros para proponer implicaciones específicas.")

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### Temas recurrentes")
            if analysis["themes"]:
                temas_df = pd.DataFrame(analysis["themes"][:6]).rename(
                    columns={"tema": "Tema", "menciones": "Menciones", "grupos": "Grupos", "mesas": "Mesas"}
                )
                st.bar_chart(temas_df.set_index("Tema")["Menciones"])
            else:
                st.caption("No hay suficiente texto codificado para identificar temas recurrentes.")
        with c2:
            st.markdown("#### Barreras recurrentes")
            if analysis["barriers"]:
                barr_df = pd.DataFrame(analysis["barriers"], columns=["Barrera", "Menciones"])
                st.bar_chart(barr_df.set_index("Barrera"))
                st.caption("Las barras representan menciones registradas, no número de participantes.")
            else:
                st.caption("No hay barreras clasificadas en los registros seleccionados.")

        if scope_label == "Sesión completa" and analysis["mesa_summaries"]:
            st.markdown("#### Lectura por mesa")
            for mesa_name, bullets_mesa in analysis["mesa_summaries"]:
                with st.expander(mesa_name):
                    for item in bullets_mesa:
                        st.markdown(f"- {item}")

        st.divider()
        st.markdown("### Edición final antes de descargar")
        st.caption(
            "Puedes ajustar el texto generado. Los archivos Word y PDF usarán exactamente esta versión editada de la síntesis y de las prioridades."
        )
        executive_edit = st.text_area(
            "Síntesis ejecutiva",
            key=exec_key,
            height=170,
        )
        priorities_edit = st.text_area(
            "Conclusiones / prioridades",
            key=pri_key,
            height=150,
        )

        theme_chart_items = [(t["tema"], t["menciones"]) for t in analysis["themes"][:6]]
        chart_theme = crear_grafica_barras(theme_chart_items, "Temas recurrentes") if theme_chart_items else None
        chart_barrier = crear_grafica_barras(analysis["barriers"][:6], "Barreras recurrentes") if analysis["barriers"] else None

        word_bytes = generar_docx_sintesis(
            analysis,
            executive_edit,
            priorities_edit,
            chart_theme=chart_theme,
            chart_barrier=chart_barrier,
        )
        pdf_bytes = generar_pdf_sintesis(
            analysis,
            executive_edit,
            priorities_edit,
            chart_theme=chart_theme,
            chart_barrier=chart_barrier,
        )

        d1, d2 = st.columns(2)
        with d1:
            if word_bytes:
                st.download_button(
                    "⬇️ Descargar síntesis en Word",
                    data=word_bytes,
                    file_name="Sintesis_Diagnostico_Economia_Circular_Jalisco.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
            else:
                st.warning("Para exportar a Word, agrega python-docx al archivo requirements.txt.")
        with d2:
            if pdf_bytes:
                st.download_button(
                    "⬇️ Descargar síntesis en PDF",
                    data=pdf_bytes,
                    file_name="Sintesis_Diagnostico_Economia_Circular_Jalisco.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                )
            else:
                st.warning("Para exportar a PDF, agrega reportlab al archivo requirements.txt.")

        st.caption(
            "Los documentos usan Montserrat cuando la fuente está disponible en el servidor o en assets; en Word se solicita Montserrat como fuente del documento. Ambos formatos incluyen logotipos disponibles en assets y paginación en el pie."
        )


st.divider()

st.markdown(
    """
    <div class="app-footer">
        COINVIERTE · Herramienta de apoyo para la sistematización
        cualitativa de la Línea Base de Economía Circular en Jalisco
    </div>
    """,
    unsafe_allow_html=True,
)
